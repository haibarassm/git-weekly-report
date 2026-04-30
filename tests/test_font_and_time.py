"""测试字体设置和时间更新功能"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt

from src.integrations.resume.document_builder import (
    _apply_style_from_template,
    DocumentBuilder,
)


def test_bullet_font_setting():
    """测试 ● 符号字体设置"""
    from docx import Document
    from src.integrations.resume.document_builder import _apply_style_from_template
    from docx.oxml.ns import qn

    doc = Document()
    p = doc.add_paragraph()

    # 添加 ● 符号
    run = p.add_run("● ")
    template_style = {
        'font_name': '微软雅黑',
        'font_size': 10.5,
        'bold': False
    }

    # 应用样式
    _apply_style_from_template(run, template_style)

    # 验证字体
    assert run.font.name == '微软雅黑', f"字体名称应为 '微软雅黑'，实际为 '{run.font.name}'"
    assert run.font.size.pt == 10.5, f"字体大小应为 10.5pt，实际为 {run.font.size.pt}pt"

    # 验证 XML 中的字体设置
    if run.element.rPr and run.element.rPr.rFonts:
        ascii_font = run.element.rPr.rFonts.get(qn('w:ascii'))
        east_asia_font = run.element.rPr.rFonts.get(qn('w:eastAsia'))
        h_ansi_font = run.element.rPr.rFonts.get(qn('w:hAnsi'))
        cs_font = run.element.rPr.rFonts.get(qn('w:cs'))

        assert ascii_font == '微软雅黑', f"ASCII 字体应为 '微软雅黑'，实际为 '{ascii_font}'"
        assert east_asia_font == '微软雅黑', f"EastAsia 字体应为 '微软雅黑'，实际为 '{east_asia_font}'"
        assert h_ansi_font == '微软雅黑', f"hAnsi 字体应为 '微软雅黑'，实际为 '{h_ansi_font}'"
        assert cs_font == '微软雅黑', f"cs 字体应为 '微软雅黑'，实际为 '{cs_font}'"

    print("✓ test_bullet_font_setting 通过")


def test_update_project_time():
    """测试更新项目经验中的"至今"时间"""
    from src.integrations.resume.document_builder import DocumentBuilder
    import shutil

    # 创建测试配置
    class TestConfig:
        def get_output_dir(self):
            return "tests/output"

    config = TestConfig()
    builder = DocumentBuilder(config)

    # 创建测试模板（包含"至今"的项目）
    tests_dir = Path("tests")
    tests_dir.mkdir(exist_ok=True)

    template_doc = Document()
    template_doc.add_heading('个人简历', level=0)

    # 添加工作经历部分
    p = template_doc.add_paragraph()
    run = p.add_run("公司名称: 测试公司")
    p = template_doc.add_paragraph()
    run = p.add_run("在职时间: 2020/01—至今")

    # 添加项目经验部分
    template_doc.add_heading('项目经验', level=2)
    p = template_doc.add_paragraph()
    run = p.add_run("2020/4—至今        测试项目")
    p = template_doc.add_paragraph()
    run = p.add_run("相关技术：Java, Spring Boot")
    p = template_doc.add_paragraph()
    run = p.add_run("产品描述：这是一个测试项目")

    template_path = tests_dir / "test_time_template.docx"
    template_doc.save(template_path)

    try:
        # 创建新项目（2025/4开始）
        new_projects = [
            {
                "id": "new_project",
                "name": "新项目",
                "period": "2025/4—至今",
                "tech_stack": "Python",
                "description": "新项目",
                "main_contributions": [],
                "bullets": [],
                "highlights": [],
                "company_id": "new_company"
            }
        ]

        # 创建工作经历
        work_experiences = [
            {
                "company_info": {
                    "id": "new_company",
                    "name": "新公司",
                    "industry": "互联网",
                    "position": "开发"
                },
                "work_experience": "● 测试职责",
                "company_period": "2025/4—至今",
                "projects": ["new_project"]
            }
        ]

        # 生成简历
        filepath = builder.build_with_template(
            new_projects=new_projects,
            work_experiences=work_experiences,
            template_path=str(template_path)
        )

        # 验证生成的文档
        result_doc = Document(filepath)
        text = "\n".join([p.text for p in result_doc.paragraphs])

        # 检查工作经历时间是否更新
        assert "2020/01—2025/3" in text, "工作经历时间应从 2020/01—至今 更新为 2020/01—2025/3"
        # 检查项目时间是否更新
        assert "2020/4—2025/3" in text, "项目时间应从 2020/4—至今 更新为 2020/4—2025/3"
        # 新公司和新项目应该仍显示"至今"
        assert "2025/4—至今" in text, "新公司和新项目应显示 2025/4—至今"

        print(f"✓ test_update_project_time 通过，文件: {filepath}")
    finally:
        # 清理测试文件
        if template_path.exists():
            template_path.unlink()


if __name__ == "__main__":
    print("运行测试用例...")
    print()

    test_bullet_font_setting()
    test_update_project_time()

    print()
    print("所有测试通过！")
