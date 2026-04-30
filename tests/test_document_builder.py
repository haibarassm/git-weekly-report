"""测试 DocumentBuilder 基本功能"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt

from src.integrations.resume.document_builder import (
    _set_font,
    _get_template_font_style,
    _apply_style_from_template,
)


def test_set_font():
    """测试字体设置功能"""
    # 创建一个测试文档
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("测试文字")

    # 调用 _set_font
    _set_font(run, font_name='微软雅黑', font_size=12, bold=True)

    # 验证
    assert run.font.name == '微软雅黑'
    assert run.font.size.pt == 12
    assert run.font.bold == True
    print("✓ test_set_font 通过")


def test_get_template_font_style():
    """测试从模板段落提取字体样式"""
    # 创建一个测试文档
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("测试文字")
    run.font.name = '宋体'
    run.font.size = Pt(10.5)  # 使用 Pt 设置

    # 调用 _get_template_font_style
    style = _get_template_font_style(p)

    # 验证
    assert style['font_name'] == '宋体'
    assert style['font_size'] == 10.5
    print("✓ test_get_template_font_style 通过")


def test_apply_style_from_template():
    """测试应用模板样式"""
    # 创建模板样式
    template_style = {
        'font_name': '宋体',
        'font_size': 10.5,
        'bold': False
    }

    # 创建一个测试文档
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("测试文字")

    # 调用 _apply_style_from_template
    _apply_style_from_template(run, template_style, bold=True)

    # 验证
    assert run.font.name == '宋体'
    assert run.font.size.pt == 10.5
    assert run.font.bold == True  # 应该被覆盖为 True
    print("✓ test_apply_style_from_template 通过")


def test_build_simple_resume():
    """测试从头生成简单简历"""
    from src.integrations.resume.document_builder import DocumentBuilder

    # 创建测试配置
    class TestConfig:
        def get_output_dir(self):
            return "tests/output"

    config = TestConfig()
    builder = DocumentBuilder(config)

    # 创建测试项目数据
    test_projects = [
        {
            "name": "测试项目",
            "period": "2025/1—2025/3",
            "tech_stack": "Python, FastAPI",
            "description": "这是一个测试项目",
            "main_contributions": [
                "模块1：开发了用户管理功能",
                "模块2：实现了订单处理"
            ],
            "bullets": [
                "完成用户管理模块",
                "实现订单处理功能"
            ],
            "highlights": []
        }
    ]

    # 生成简历
    try:
        filepath = builder.build(test_projects)
        assert filepath.exists()

        # 验证生成的文档可以打开
        doc = Document(filepath)
        assert len(doc.paragraphs) > 0
        print(f"✓ test_build_simple_resume 通过，文件: {filepath}")
    except Exception as e:
        pytest.fail(f"test_build_simple_resume 失败: {e}")


def test_build_with_template():
    """测试基于模板生成简历"""
    from src.integrations.resume.document_builder import DocumentBuilder
    import shutil

    # 创建测试配置
    class TestConfig:
        def get_output_dir(self):
            return "tests/output"

    config = TestConfig()
    builder = DocumentBuilder(config)

    # 准备测试模板（复制一个简单的模板）
    tests_dir = Path("tests")
    tests_dir.mkdir(exist_ok=True)

    # 创建一个简单的测试模板
    template_doc = Document()
    template_doc.add_heading('个人简历', level=0)

    # 添加工作经历部分
    p = template_doc.add_paragraph()
    run = p.add_run("公司名称: 测试公司")
    p = template_doc.add_paragraph()
    run = p.add_run("在职时间: 2020/01—至今")
    p = template_doc.add_paragraph()
    run = p.add_run("所属行业：互联网")
    p = template_doc.add_paragraph()
    run = p.add_run("主要职责：")
    p = template_doc.add_paragraph()
    run = p.add_run("● 负责开发工作")

    # 添加项目经验部分
    template_doc.add_heading('项目经验', level=2)

    template_path = tests_dir / "test_template.docx"
    template_doc.save(template_path)

    # 创建新项目数据
    new_projects = [
        {
            "id": "test_project",
            "name": "新测试项目",
            "period": "2025/4—至今",
            "tech_stack": "Java, Spring Boot",
            "description": "这是一个新的测试项目",
            "main_contributions": [
                "支付订单：负责订单创建流程",
                "退款：实现退款功能"
            ],
            "bullets": [
                "完成支付订单模块",
                "实现退款功能"
            ],
            "highlights": ["对接支付渠道"],
            "company_id": "test_company"
        }
    ]

    # 创建工作经历数据
    work_experiences = [
        {
            "company_info": {
                "id": "test_company",
                "name": "杭州碰呗网络科技有限公司",
                "industry": "互联网",
                "position": "java开发工程师"
            },
            "work_experience": "● 支付订单：负责订单创建流程\n● 退款：实现退款功能",
            "company_period": "2025/4—至今",
            "projects": ["test_project"]
        }
    ]

    try:
        # 生成简历
        filepath = builder.build_with_template(
            new_projects=new_projects,
            work_experiences=work_experiences,
            template_path=str(template_path)
        )

        assert filepath.exists()

        # 验证生成的文档
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])

        # 检查是否包含新公司和项目
        assert "杭州碰呗网络科技有限公司" in text
        assert "新测试项目" in text
        # 检查是否更新了上一家公司的时间
        assert "2020/01—2025/3" in text

        print(f"✓ test_build_with_template 通过，文件: {filepath}")
    except Exception as e:
        pytest.fail(f"test_build_with_template 失败: {e}")
    finally:
        # 清理测试文件
        if template_path.exists():
            template_path.unlink()


if __name__ == "__main__":
    print("运行测试用例...")
    print()

    test_set_font()
    test_get_template_font_style()
    test_apply_style_from_template()
    test_build_simple_resume()
    test_build_with_template()

    print()
    print("所有测试通过！")

