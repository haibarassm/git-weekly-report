"""Word 文档构建器 - 生成简历 Word 文档"""
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional
import logging
import re
import copy

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx 未安装，请运行: pip install python-docx")


def _copy_paragraph_style(source_para, target_para):
    """复制段落的所有样式（包括段落样式和 run 样式）"""
    # 复制段落样式
    if source_para.style:
        target_para.style = source_para.style

    # 复制段落格式
    if source_para.alignment:
        target_para.alignment = source_para.alignment
    if source_para.paragraph_format.left_indent:
        target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    if source_para.paragraph_format.right_indent:
        target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
    if source_para.paragraph_format.first_line_indent:
        target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
    if source_para.paragraph_format.line_spacing:
        target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    if source_para.paragraph_format.space_before:
        target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    if source_para.paragraph_format.space_after:
        target_para.paragraph_format.space_after = source_para.paragraph_format.space_after


def _add_text_with_style_from_template(doc, text, template_para, bold=None):
    """添加文本，从模板段落复制样式

    Args:
        doc: 文档对象
        text: 要添加的文本
        template_para: 模板段落
        bold: 是否加粗（如果指定，覆盖模板的加粗设置）

    Returns:
        新创建的 run 对象
    """
    p = doc.add_paragraph()

    # 复制段落样式
    _copy_paragraph_style(template_para, p)

    # 添加文本
    run = p.add_run(text)

    # 如果模板段落有 run，复制第一个 run 的样式
    if template_para.runs:
        for template_run in template_para.runs:
            if template_run.text.strip():
                # 复制字体属性
                if template_run.font.name:
                    run.font.name = template_run.font.name
                if template_run.font.size:
                    run.font.size = template_run.font.size
                if bold is not None:
                    run.font.bold = bold
                elif template_run.font.bold is not None:
                    run.font.bold = template_run.font.bold

                # 复制中文字体
                try:
                    if template_run.element.rPr is not None and template_run.element.rPr.rFonts is not None:
                        east_asia = template_run.element.rPr.rFonts.get(qn('w:eastAsia'))
                        if east_asia and run.element.rPr is not None:
                            if run.element.rPr.rFonts is not None:
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
                            else:
                                # 创建 rFonts 元素
                                from docx.oxml import OxmlElement
                                rFonts = OxmlElement('w:rFonts')
                                rFonts.set(qn('w:eastAsia'), east_asia)
                                run.element.rPr.append(rFonts)
                except (AttributeError, TypeError):
                    pass
                break

    return p, run


def _set_font(run, font_name='微软雅黑', font_size=10.5, bold=False):
    """设置字体样式（仅在无法从模板获取样式时使用）"""
    if run.font.name is None:  # 只在没有设置字体时才设置
        run.font.name = font_name
    if run.font.size is None:
        run.font.size = Pt(font_size)
    if bold is not None and run.font.bold is None:
        run.font.bold = bold
    # 设置中文字体（需要检查 rPr 是否存在）
    try:
        if run.element.rPr is not None and run.element.rPr.rFonts is not None:
            if run.element.rPr.rFonts.get(qn('w:eastAsia')) is None:
                run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except (AttributeError, TypeError):
        pass


def _get_template_font_style(template_para):
    """从模板段落中提取字体样式

    Returns:
        dict: {'font_name', 'font_size', 'bold'}
    """
    style = {
        'font_name': '微软雅黑',
        'font_size': 10.5,
        'bold': False
    }

    if not template_para or not template_para.runs:
        return style

    # 从第一个有文字的 run 中获取样式
    for run in template_para.runs:
        if run.text.strip():
            if run.font.name:
                style['font_name'] = run.font.name
            if run.font.size:
                style['font_size'] = run.font.size.pt
            if run.font.bold is not None:
                style['bold'] = run.font.bold

            # 尝试获取中文字体（需要检查 rPr 是否存在）
            try:
                if run.element.rPr is not None and run.element.rPr.rFonts is not None:
                    east_asia = run.element.rPr.rFonts.get(qn('w:eastAsia'))
                    if east_asia:
                        style['font_name'] = east_asia
            except (AttributeError, TypeError):
                pass
            break

    return style


def _apply_style_from_template(run, template_style, bold=None):
    """应用模板样式到 run

    Args:
        run: 要应用样式的 run 对象
        template_style: 从 _get_template_font_style 获取的样式字典
        bold: 是否加粗（覆盖模板中的加粗设置）
    """
    if not DOCX_AVAILABLE:
        return

    from docx.oxml import OxmlElement

    font_name = template_style.get('font_name', '微软雅黑')
    font_size = template_style.get('font_size', 10.5)
    template_bold = template_style.get('bold', False)

    # 设置基本字体属性
    run.font.name = font_name
    run.font.size = Pt(font_size)
    # 如果明确指定了 bold，使用指定的值；否则使用模板中的值
    if bold is not None:
        run.font.bold = bold
    else:
        run.font.bold = template_bold

    # 确保 rPr 元素存在
    if run.element.rPr is None:
        rPr = OxmlElement('w:rPr')
        run.element.insert(0, rPr)
    else:
        rPr = run.element.rPr

    # 设置或创建 rFonts 元素
    if rPr.rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    else:
        rFonts = rPr.rFonts

    # 设置所有字体类型以确保一致性
    rFonts.set(qn('w:ascii'), font_name)      # ASCII 字体（英文、符号）
    rFonts.set(qn('w:eastAsia'), font_name)   # East Asian 字体（中文）
    rFonts.set(qn('w:hAnsi'), font_name)      # High ANSI 字体
    rFonts.set(qn('w:cs'), font_name)         # Complex Script 字体


def _add_paragraph_with_markdown(doc, text: str, style=None, font_size=11):
    """添加段落，支持 Markdown 加粗格式 **text**，清理 • 符号"""
    p = doc.add_paragraph(style=style)

    # 清理 • 符号（列表标记）
    text = text.lstrip('•').strip()

    # 解析 **text** 格式
    parts = re.split(r'(\*\*.+?\*\*)', text)

    for part in parts:
        if not part:
            continue

        # 检查是否是加粗部分
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            if content:
                run = p.add_run(content)
                _set_font(run, font_size=font_size, bold=True)
        else:
            if part.strip() or not p.runs:
                run = p.add_run(part)
                _set_font(run, font_size=font_size)

    return p


class DocumentBuilder:
    """简历 Word 文档构建器"""

    def __init__(self, config):
        self.config = config
        self.output_dir = Path(config.get_output_dir())
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def build(self, resume_projects: List[Dict]) -> Path:
        """构建 Word 文档（从头生成）"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx 未安装，请运行: pip install python-docx")

        logger = logging.getLogger(__name__)
        doc = Document()

        # 标题
        title = doc.add_heading('Resume', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 项目列表
        for project_idx, project in enumerate(resume_projects):
            # 第一行：时间 + 项目名称（用tab分隔）
            p = doc.add_paragraph()
            period = project.get("period", "")
            name = project.get("name", "")
            run = p.add_run(f"{period}\t{name}")
            _set_font(run)

            # 第二行：相关技术
            if project.get("tech_stack"):
                p = doc.add_paragraph()
                run = p.add_run("相关技术：")
                _set_font(run, bold=True)
                run = p.add_run(project["tech_stack"])
                _set_font(run)

            # 第三行：产品描述
            if project.get("description"):
                p = doc.add_paragraph()
                run = p.add_run("产品描述：")
                _set_font(run, bold=True)
                run = p.add_run(project["description"])
                _set_font(run)

            # 第四行：主要贡献
            if project.get("main_contributions"):
                p = doc.add_paragraph()
                run = p.add_run("主要贡献：")
                _set_font(run, bold=True)

                for contribution in project["main_contributions"]:
                    p = doc.add_paragraph()
                    run = p.add_run("● ")
                    _set_font(run)
                    run = p.add_run(contribution)
                    _set_font(run)

            # 第五行：关键成果
            if project.get("bullets"):
                p = doc.add_paragraph()
                run = p.add_run("关键成果：")
                _set_font(run, bold=True)

                for bullet in project.get("bullets", []):
                    p = doc.add_paragraph()
                    clean_bullet = bullet.lstrip('●').lstrip('•').strip()
                    run = p.add_run("● ")
                    _set_font(run)
                    parts = re.split(r'(\*\*.+?\*\*)', clean_bullet)
                    for part in parts:
                        if not part:
                            continue
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            _set_font(run, bold=True)
                        else:
                            run = p.add_run(part)
                            _set_font(run)

            # 项目亮点
            if project.get("highlights"):
                p = doc.add_paragraph()
                run = p.add_run("项目亮点：")
                _set_font(run, bold=True)

                for highlight in project["highlights"]:
                    p = doc.add_paragraph()
                    run = p.add_run("● ")
                    _set_font(run)
                    run = p.add_run(highlight)
                    _set_font(run)

            # 分隔线
            if project_idx < len(resume_projects) - 1:  # 不是最后一个项目
                p = doc.add_paragraph()
                run = p.add_run("----------------------------------------------------------------------")
                _set_font(run)

            doc.add_paragraph()

        # 保存文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_{timestamp}.docx"
        filepath = self.output_dir / filename
        doc.save(filepath)
        logger.info(f">>> [文档构建] 保存: {filepath}")
        return filepath

    def build_with_template(self, new_projects: List[Dict], work_experiences: List[Dict], template_path: str) -> Path:
        """基于历史简历模板生成新文档（插入工作经历 + 项目经验）

        Args:
            new_projects: 新生成的项目列表
            work_experiences: 公司工作经历列表
            template_path: 历史简历文件路径

        Returns:
            生成的文件路径
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx 未安装，请运行: pip install python-docx")

        logger = logging.getLogger(__name__)

        # 打开模板文档
        doc = Document(template_path)

        # 自动更新上一家公司的"至今"时间
        self._update_previous_company_end_date(doc, new_projects)

        # 从模板中提取常用样式
        template_styles = self._extract_template_styles(doc)
        logger.info(f">>> [文档构建] 从模板提取了 {len(template_styles)} 种样式")

        # 第一步：插入工作经历（如果有）
        if work_experiences:
            work_insert_index = self._find_work_experience_insert_index(doc)
            logger.info(f">>> [文档构建] 工作经历插入位置: {work_insert_index}")

            # 按公司插入工作经历（倒序，最新的公司在最前）
            for idx, work_exp in enumerate(reversed(work_experiences)):
                # 插入公司信息
                company_info = work_exp["company_info"]

                # 公司名称
                p = doc.add_paragraph()
                run = p.add_run("公司名称: ")
                _apply_style_from_template(run, template_styles['label'], bold=True)
                run = p.add_run(company_info["name"])
                _apply_style_from_template(run, template_styles['content'])
                if work_insert_index < len(doc.paragraphs):
                    doc.paragraphs[work_insert_index]._element.addprevious(p._element)
                else:
                    doc.element.body.append(p._element)
                work_insert_index += 1

                # 在职时间
                p = doc.add_paragraph()
                run = p.add_run("在职时间: ")
                _apply_style_from_template(run, template_styles['label'], bold=True)
                run = p.add_run(work_exp["company_period"])
                _apply_style_from_template(run, template_styles['content'])
                if work_insert_index < len(doc.paragraphs):
                    doc.paragraphs[work_insert_index]._element.addprevious(p._element)
                else:
                    doc.element.body.append(p._element)
                work_insert_index += 1

                # 所属行业
                if company_info.get("industry"):
                    p = doc.add_paragraph()
                    run = p.add_run("所属行业：")
                    _apply_style_from_template(run, template_styles['label'], bold=True)
                    run = p.add_run(company_info["industry"])
                    _apply_style_from_template(run, template_styles['content'])
                    if work_insert_index < len(doc.paragraphs):
                        doc.paragraphs[work_insert_index]._element.addprevious(p._element)
                    else:
                        doc.element.body.append(p._element)
                    work_insert_index += 1

                # 主要职责
                p = doc.add_paragraph()
                run = p.add_run("主要职责：")
                _apply_style_from_template(run, template_styles['label'], bold=True)
                if work_insert_index < len(doc.paragraphs):
                    doc.paragraphs[work_insert_index]._element.addprevious(p._element)
                else:
                    doc.element.body.append(p._element)
                work_insert_index += 1

                # 插入职责描述（每行一条）
                for line in work_exp["work_experience"].split('\n'):
                    if not line.strip():
                        continue
                    p = doc.add_paragraph()

                    # 解析 ● 符号和内容
                    line = line.strip()
                    if line.startswith('●'):
                        # 分离 ● 符号和内容
                        parts = line.split('：', 1) if '：' in line else [line, '']
                        if len(parts) == 2:
                            # 格式：● 模块名：内容
                            run = p.add_run(parts[0] + '：')
                            _apply_style_from_template(run, template_styles['bullet'])
                            run = p.add_run(parts[1])
                            _apply_style_from_template(run, template_styles['content'])
                        else:
                            # 格式：● 内容
                            run = p.add_run("● ")
                            _apply_style_from_template(run, template_styles['bullet'])
                            run = p.add_run(line[1:].strip())  # 去掉 ● 符号
                            _apply_style_from_template(run, template_styles['content'])
                    else:
                        # 没有 ● 符号，添加一个
                        run = p.add_run("● ")
                        _apply_style_from_template(run, template_styles['bullet'])
                        run = p.add_run(line.strip())
                        _apply_style_from_template(run, template_styles['content'])

                    if work_insert_index < len(doc.paragraphs):
                        doc.paragraphs[work_insert_index]._element.addprevious(p._element)
                    else:
                        doc.element.body.append(p._element)
                    work_insert_index += 1

                # 公司之间用分隔线（包括最后一个公司）
                p = doc.add_paragraph()
                run = p.add_run("----------------------------------------------------------------------")
                _set_font(run)
                if work_insert_index < len(doc.paragraphs):
                    doc.paragraphs[work_insert_index]._element.addprevious(p._element)
                else:
                    doc.element.body.append(p._element)
                work_insert_index += 1

                # 添加公司后的空行
                p = doc.add_paragraph()
                if work_insert_index < len(doc.paragraphs):
                    doc.paragraphs[work_insert_index]._element.addprevious(p._element)
                else:
                    doc.element.body.append(p._element)
                work_insert_index += 1

            logger.info(f">>> [文档构建] 已插入 {len(work_experiences)} 个公司工作经历")

        # 第二步：插入项目经验（保持原有逻辑）
        insert_index = self._find_project_section_insert_index(doc)
        logger.info(f">>> [文档构建] 项目经验插入位置: {insert_index}")

        # 找到项目经验部分的标题（插入点）
        insert_index = self._find_project_section_insert_index(doc)
        logger.info(f">>> [文档构建] 项目经验插入位置: {insert_index}")

        # 如果没找到项目经验标题，返回 None 时不添加
        if insert_index is None:
            logger.warning(">>> [文档构建] 未找到项目经验标题，新项目将添加到文档末尾")
            insert_index = len(doc.paragraphs)  # 设置为文档末尾
            # 添加项目经验标题
            heading = doc.add_heading("项目经验", level=2)
            doc.element.body.insert(insert_index, heading._element)
            insert_index += 1
            # 添加空行
            p = doc.add_paragraph()
            doc.element.body.insert(insert_index, p._element)
            insert_index += 1

        # 在项目经验标题后插入新项目（倒序，最新的在最前面）
        logger.info(f">>> [文档构建] 开始插入 {len(new_projects)} 个新项目，起始位置: {insert_index}")

        # 获取"项目经验"段落的 _element 引用，用于在其后插入
        if insert_index > 0 and insert_index <= len(doc.paragraphs):
            # 使用段落对象而不是索引来插入
            insert_before_para = doc.paragraphs[insert_index] if insert_index < len(doc.paragraphs) else None
            logger.info(f">>> [文档构建] 将在段落 '{insert_before_para.text[:30] if insert_before_para else '文档末尾'}' 之前插入")
        else:
            insert_before_para = None

        for project_idx, project in enumerate(reversed(new_projects)):
            logger.info(f">>> [文档构建] 插入项目 {project_idx + 1}/{len(new_projects)}: {project['name']}")

            # 第一行：时间 + 项目名称（用tab分隔）
            p = doc.add_paragraph()
            period = project.get("period", "")
            name = project.get("name", "")
            run = p.add_run(f"{period}\t{name}")
            _set_font(run)
            if insert_before_para:
                insert_before_para._element.addprevious(p._element)
            else:
                doc.element.body.append(p._element)
            logger.info(f">>> [文档构建]   插入项目标题行: {period} {name}")

            # 第二行：相关技术
            if project.get("tech_stack"):
                p = doc.add_paragraph()
                run = p.add_run("相关技术：")
                _set_font(run, bold=True)
                run = p.add_run(project["tech_stack"])
                _set_font(run)
                if insert_before_para:
                    insert_before_para._element.addprevious(p._element)
                else:
                    doc.element.body.append(p._element)
                logger.info(f">>> [文档构建]   插入相关技术")

            # 第三行：产品描述
            if project.get("description"):
                p = doc.add_paragraph()
                run = p.add_run("产品描述：")
                _set_font(run, bold=True)
                run = p.add_run(project["description"])
                _set_font(run)
                if insert_before_para:
                    insert_before_para._element.addprevious(p._element)
                else:
                    doc.element.body.append(p._element)
                logger.info(f">>> [文档构建]   插入产品描述")

            # 第四行：主要贡献（从 summary.main_contributions 获取）
            if project.get("main_contributions"):
                p = doc.add_paragraph()
                run = p.add_run("主要贡献：")
                _set_font(run, bold=True)
                if insert_before_para:
                    insert_before_para._element.addprevious(p._element)
                else:
                    doc.element.body.append(p._element)
                logger.info(f">>> [文档构建]   插入主要贡献")

                for contribution in project["main_contributions"]:
                    p = doc.add_paragraph()
                    run = p.add_run("● ")
                    _set_font(run)
                    run = p.add_run(contribution)
                    _set_font(run)
                    if insert_before_para:
                        insert_before_para._element.addprevious(p._element)
                    else:
                        doc.element.body.append(p._element)

            # 第五行：关键成果（bullets）
            if project.get("bullets"):
                p = doc.add_paragraph()
                run = p.add_run("关键成果：")
                _set_font(run, bold=True)
                if insert_before_para:
                    insert_before_para._element.addprevious(p._element)
                else:
                    doc.element.body.append(p._element)
                logger.info(f">>> [文档构建]   插入关键成果")

                # 插入 bullets（每个bullet一行）
                for bullet in project.get("bullets", []):
                    p = doc.add_paragraph()
                    # 清理 • 符号并添加内容
                    clean_bullet = bullet.lstrip('●').lstrip('•').strip()
                    # 添加列表符号
                    run = p.add_run("● ")
                    _apply_style_from_template(run, template_styles['bullet'])
                    # 解析 **加粗** 格式
                    parts = re.split(r'(\*\*.+?\*\*)', clean_bullet)
                    for part in parts:
                        if not part:
                            continue
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            _apply_style_from_template(run, template_styles['content'], bold=True)
                        else:
                            run = p.add_run(part)
                            _apply_style_from_template(run, template_styles['content'])
                    if insert_before_para:
                        insert_before_para._element.addprevious(p._element)
                    else:
                        doc.element.body.append(p._element)

            # 插入项目亮点（如果有）
            if project.get("highlights"):
                p = doc.add_paragraph()
                run = p.add_run("项目亮点：")
                _set_font(run, bold=True)
                if insert_before_para:
                    insert_before_para._element.addprevious(p._element)
                else:
                    doc.element.body.append(p._element)
                logger.info(f">>> [文档构建]   插入项目亮点")

                for highlight in project["highlights"]:
                    p = doc.add_paragraph()
                    run = p.add_run("● ")
                    _set_font(run)
                    run = p.add_run(highlight)
                    _set_font(run)
                    if insert_before_para:
                        insert_before_para._element.addprevious(p._element)
                    else:
                        doc.element.body.append(p._element)

            # 项目之间用分隔线（包括最后一个项目）
            p = doc.add_paragraph()
            run = p.add_run("----------------------------------------------------------------------")
            _apply_style_from_template(run, template_styles['content'])
            if insert_before_para:
                insert_before_para._element.addprevious(p._element)
            else:
                doc.element.body.append(p._element)
            logger.info(f">>> [文档构建]   插入项目分隔线")

        # 保存文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_updated_{timestamp}.docx"
        filepath = self.output_dir / filename
        doc.save(filepath)
        logger.info(f">>> [文档构建] 保存更新简历: {filepath}")
        return filepath

    def _find_project_section_insert_index(self, doc: Document) -> Optional[int]:
        """查找项目经验部分的插入位置（标题后的第一个位置）

        策略：
        1. 查找"工作经历"、"项目经验"、"实习经历"等章节标题（不限定样式）
        2. 如果找到，返回标题后的第一个非空位置
        3. 如果找不到，返回 None（会在文档末尾创建新章节）
        """
        logger = logging.getLogger(__name__)

        # 先打印所有段落帮助调试
        logger.info(f">>> [查找插入位置] 文档共有 {len(doc.paragraphs)} 个段落")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:  # 只打印非空段落
                logger.info(f"  段落 {i}: '{text[:50]}' (样式: {para.style.name})")
            elif text == '':  # 空行也打印
                logger.info(f"  段落 {i}: [空行] (样式: {para.style.name})")

        # 优先精确匹配"项目经验"相关标题（按优先级排序）
        # 注意：需要排除"技术技能"、"专业技能"等干扰项
        project_keywords = ['项目经验', '项目介绍', '项目展示', 'Project Experience']
        work_keywords = ['工作经历', '工作', 'Work Experience', '实习经历']
        exclude_keywords = ['技术技能', '专业技能', '技能', 'Skills', '教育', 'Education']

        # 第一优先级：精确匹配"项目经验"章节（排除技术技能等干扰）
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # 检查是否包含排除关键词
            if any(exclude in text for exclude in exclude_keywords):
                logger.info(f">>> [查找插入位置] 跳过干扰项: '{text}' 在位置 {i}")
                continue
            # 精确匹配项目经验关键词
            for keyword in project_keywords:
                if keyword in text:
                    logger.info(f">>> [查找插入位置] ✓✓✓ 找到项目经验标题: '{text}' 在位置 {i}")
                    return self._find_insert_position_after_header(doc, i)

        # 第二优先级：查找工作经历章节（在其内部找项目部分）
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # 检查是否包含排除关键词
            if any(exclude in text for exclude in exclude_keywords):
                continue
            for keyword in work_keywords:
                if keyword in text:
                    logger.info(f">>> [查找插入位置] 找到工作经历标题: '{text}' 在位置 {i}")
                    # 在工作经历后查找是否已有项目内容，或在其后插入
                    return self._find_insert_position_after_header(doc, i)

        # 第三优先级：查找任何包含"项目"的标题（但要排除"工作经历"和技术技能）
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # 检查是否包含排除关键词
            if any(exclude in text for exclude in exclude_keywords + ['工作', '实习']):
                continue
            if '项目' in text:
                logger.info(f">>> [查找插入位置] 找到项目相关标题: '{text}' 在位置 {i}")
                return self._find_insert_position_after_header(doc, i)

        # 如果完全找不到，返回 None（会创建新的项目经验标题）
        logger.info(f">>> [查找插入位置] 未找到项目经验标题，返回 None")
        return None

    def _find_insert_position_after_header(self, doc: Document, header_index: int) -> int:
        """找到章节标题后的实际内容插入位置

        跳过空行和装饰线，找到第一个实际内容的位置
        """
        logger = logging.getLogger(__name__)

        insert_pos = header_index + 1
        while insert_pos < len(doc.paragraphs):
            next_text = doc.paragraphs[insert_pos].text.strip()
            # 如果是空行或纯装饰线（如连续的 - 或 =），继续往后找
            if not next_text or next_text in ['-', '=', '―', '—', '____', '______________']:
                logger.info(f">>> [查找插入位置] 跳过空行/装饰线: 位置 {insert_pos} 内容: '{next_text}'")
                insert_pos += 1
            else:
                # 找到实际内容位置 - 检查是否是另一个章节标题
                # 如果下一个是章节标题（如"工作经历"、"教育"等），应该在这个位置插入
                # 如果下一个是项目内容，也应该在这个位置插入（在现有项目之前）
                logger.info(f">>> [查找插入位置] 找到插入位置: {insert_pos}")
                logger.info(f">>> [查找插入位置] 标题位置: {header_index} ('{doc.paragraphs[header_index].text.strip()}')")
                logger.info(f">>> [查找插入位置] 插入点内容: '{next_text[:50]}'")
                return insert_pos

        # 如果后面没有内容了，就在文档末尾插入
        logger.info(f">>> [查找插入位置] 章节后无内容，插入到文档末尾: {len(doc.paragraphs)}")
        return len(doc.paragraphs)

    def _update_previous_company_end_date(self, doc: Document, new_projects: List[Dict]):
        """自动更新上一家公司的"至今"结束时间

        策略：
        1. 找到模板中第一个"至今"的公司
        2. 计算新项目的最早开始时间
        3. 把"至今"改为新项目开始时间的前一个月

        Args:
            doc: Word 文档
            new_projects: 新项目列表
        """
        logger = logging.getLogger(__name__)

        # 计算新项目的最早开始时间
        earliest_start = None
        for project in new_projects:
            period = project.get("period", "")
            if not period:
                continue

            # 解析时间范围（支持中文和英文连字符）
            parts = period.replace("—", "-").split("-")
            if len(parts) >= 1:
                start_part = parts[0].strip()
                try:
                    if earliest_start is None or start_part < earliest_start:
                        earliest_start = start_part
                except:
                    pass

        if not earliest_start:
            logger.info(">>> [自动更新时间] 无新项目时间信息，跳过更新")
            return

        # 计算结束时间（新项目开始时间的前一个月）
        try:
            from datetime import datetime, timedelta
            start_date = datetime.strptime(earliest_start, "%Y/%m")
            end_date = start_date - timedelta(days=1)
            end_str = f"{end_date.year}/{end_date.month}"
            logger.info(f">>> [自动更新时间] 新项目开始: {earliest_start}, 上一家公司结束: {end_str}")
        except Exception as e:
            logger.warning(f">>> [自动更新时间] 时间计算失败: {e}")
            return

        # 查找并更新第一个"至今"的公司
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # 查找包含"至今"的在职时间行
            if ("在职时间:" in text or "在职时间：" in text) and "至今" in text:
                # 检查是否是第一个"至今"的公司（通常在模板前面）
                # 保留原始格式，只替换"至今"文本
                updated = False
                for run in para.runs:
                    if "至今" in run.text:
                        run.text = run.text.replace("至今", end_str)
                        updated = True
                        logger.info(f">>> [自动更新时间] 已更新 run: '{run.text}'")
                        break
                if updated:
                    logger.info(f">>> [自动更新时间] 已更新段落: '{text}' -> '{para.text.strip()}'")
                    break  # 只更新第一个"至今"的公司

        # 查找并更新所有项目经验中的"至今"
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # 项目经验的时间格式通常是：YYYY/M—至今    项目名称
            # 检查是否包含"至今"且不是工作经历部分
            if "至今" in text and ("在职时间:" not in text and "在职时间：" not in text):
                # 检查是否是项目时间行（通常以时间开头）
                if text[0].isdigit() or text.startswith(" "):
                    # 保留原始格式，只替换"至今"文本
                    for run in para.runs:
                        if "至今" in run.text:
                            run.text = run.text.replace("至今", end_str)
                            logger.info(f">>> [自动更新时间] 已更新项目时间: 段{i} '{run.text}'")
                            break

    def _extract_template_styles(self, doc: Document) -> Dict:
        """从模板文档中提取常用样式

        Returns:
            dict: 样式字典，包含各种内容类型的字体样式
        """
        logger = logging.getLogger(__name__)

        # 先尝试从模板的 Normal 样式中提取默认值
        default_style = {'font_name': '微软雅黑', 'font_size': 10.5, 'bold': False}
        try:
            if doc.styles and 'Normal' in doc.styles:
                normal_style = doc.styles['Normal']
                if normal_style.font.name:
                    default_style['font_name'] = normal_style.font.name
                if normal_style.font.size:
                    default_style['font_size'] = normal_style.font.size.pt
        except:
            pass

        styles = {
            'normal': default_style,  # 默认样式
            'label': {},  # 标签样式（如"公司名称:"、"在职时间:"）
            'content': {},  # 内容样式
            'bullet': {},  # 列表项样式
        }

        # 遍历模板，提取各种样式
        for para in doc.paragraphs:
            text = para.text.strip()

            # 提取标签样式（"公司名称:", "在职时间:", "相关技术:" 等）
            if any(label in text for label in ['公司名称:', '在职时间:', '相关技术:', '产品描述:', '主要贡献:', '关键成果:', '项目亮点:', '所属行业：', '主要职责：']):
                if para.runs:
                    for run in para.runs:
                        if run.text.strip() and ':' in run.text:
                            styles['label'] = _get_template_font_style(para)
                            break

            # 提取内容样式（标签后的普通文本）
            elif any(label in text for label in ['杭州碰呗网络科技有限公司', '2025/4—至今', '移动互联网']):
                if para.runs:
                    styles['content'] = _get_template_font_style(para)
                    break

            # 提取列表项样式（以 ● 开头的行）
            elif text.startswith('●'):
                styles['bullet'] = _get_template_font_style(para)
                break

            # 如果已经提取了所有样式，停止
            if all(styles.values()):
                break

        # 如果没有找到特定样式，使用默认值
        if not styles['label']:
            styles['label'] = styles['normal']
        if not styles['content']:
            styles['content'] = styles['normal']
        if not styles['bullet']:
            styles['bullet'] = styles['normal']

        logger.info(f">>> [样式提取] 标签样式: {styles['label']}")
        logger.info(f">>> [样式提取] 内容样式: {styles['content']}")
        logger.info(f">>> [样式提取] 列表样式: {styles['bullet']}")

        return styles

    def _find_work_experience_insert_index(self, doc: Document) -> int:
        """查找工作经历部分的插入位置

        Args:
            doc: Word 文档

        Returns:
            插入位置索引
        """
        logger = logging.getLogger(__name__)

        # 先打印所有段落帮助调试
        logger.info(f">>> [查找工作经历位置] 文档共有 {len(doc.paragraphs)} 个段落")

        # 查找"工作经历"标题
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            if '工作经历' in text:
                logger.info(f">>> [查找工作经历位置] 找到工作经历标题: '{text}' 在位置 {i}")
                # 返回标题后的下一个位置
                return i + 1

        # 如果没找到，返回文档末尾
        logger.info(f">>> [查找工作经历位置] 未找到工作经历标题，返回文档末尾")
        return len(doc.paragraphs)
